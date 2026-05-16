"""Tygodniowy raport DOCX z wynikami matchingu."""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import _Cell

from src.matching.compare import JDParsed
from src.matching.pipeline import MatchedOffer

_DARK_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
_GRAY = RGBColor(0x94, 0xA3, 0xB8)
_BLACK = RGBColor(0x1E, 0x29, 0x3B)

_FONT_NAME = "Calibri"
_FONT_SIZE_BODY = Pt(10)
_FONT_SIZE_H1 = Pt(16)
_FONT_SIZE_H2 = Pt(13)
_FONT_SIZE_H3 = Pt(11)

TOP_N = 10


def _set_cell_bg(cell: _Cell, hex_color: str) -> None:
    """Ustawia kolor tla komorki (niskopoziomowy XML — nazwy zgodne z OOXML)."""
    tc = cell._tc  # noqa: SLF001
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_hyperlink(paragraph: Any, url: str, text: str) -> None:
    """Dodaje klikalny hyperlink do paragrafu DOCX."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    run_elem.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text
    run_elem.append(t)
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _h1(doc: Any, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = _FONT_NAME
    run.font.size = _FONT_SIZE_H1
    run.font.bold = True
    run.font.color.rgb = _DARK_BLUE


def _h2(doc: Any, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = _FONT_NAME
    run.font.size = _FONT_SIZE_H2
    run.font.bold = True
    run.font.color.rgb = _DARK_BLUE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def _body(doc: Any, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = _FONT_NAME
    run.font.size = _FONT_SIZE_BODY
    run.font.bold = bold
    run.font.color.rgb = _BLACK


def _add_summary_section(
    doc: Any,
    matched: list[MatchedOffer],
    our_jd: JDParsed | None,
) -> None:
    _h2(doc, "1. Podsumowanie")
    scored = [m for m in matched if m.match_score is not None]
    _body(doc, f"Laczna liczba ofert po deduplikacji: {len(matched)}")
    if our_jd is not None:
        _body(doc, f"Oferty z obliczonym score (vs nasza oferta): {len(scored)}")

    doc.add_paragraph()
    portal_counts: dict[str, int] = {}
    for offer in matched:
        portal = offer.deduped.primary.portal
        portal_counts[portal] = portal_counts.get(portal, 0) + 1

    if portal_counts:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        labels = ["Portal", "Liczba ofert"]
        for cell, label in zip(hdr, labels, strict=False):
            cell.text = label
            _set_cell_bg(cell, "1F3A5F")
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = _FONT_NAME
            run.font.size = _FONT_SIZE_BODY

        for portal, count in sorted(portal_counts.items()):
            row = table.add_row().cells
            for cell, value in zip(row, [portal, str(count)], strict=False):
                cell.text = value
                for run in cell.paragraphs[0].runs:
                    run.font.name = _FONT_NAME
                    run.font.size = _FONT_SIZE_BODY


def _add_top10_section(
    doc: Any,
    matched: list[MatchedOffer],
    our_jd: JDParsed | None,
) -> None:
    _h2(doc, f"2. TOP {TOP_N} najlepszych dopasowan")

    sorted_offers = sorted(
        matched,
        key=lambda m: m.match_score.total if m.match_score else -1,
        reverse=True,
    )
    top = sorted_offers[:TOP_N]

    if not top:
        _body(doc, "Brak ofert spelniajacych kryteria.")
        return

    for i, offer in enumerate(top, start=1):
        p = offer.deduped.primary
        ms = offer.match_score
        score_str = f"{ms.total}/100" if ms else "brak score"

        heading_p = doc.add_paragraph()
        run = heading_p.add_run(f"{i}. {p.title} — {p.company}  (Score: {score_str})")
        run.font.name = _FONT_NAME
        run.font.size = _FONT_SIZE_H3
        run.font.bold = True
        run.font.color.rgb = _DARK_BLUE
        heading_p.paragraph_format.space_before = Pt(8)

        sal = p.salary
        sal_str = f"{sal.min}-{sal.max} {sal.currency}/{sal.period}" if sal else "brak stawki"
        _body(doc, f"   {p.portal} · {p.location or '—'} · {p.work_mode or '—'} · {sal_str}")

        if p.tech_stack:
            _body(doc, f"   Tech: {', '.join(p.tech_stack)}")

        link_p = doc.add_paragraph("   ")
        link_p.add_run("Link: ").font.name = _FONT_NAME
        _add_hyperlink(link_p, p.url, p.url)

        if our_jd is not None and ms is not None:
            diffs: list[str] = []
            if not ms.seniority_match:
                diffs.append(f"seniority: nasza={our_jd.seniority}, ich={p.seniority}")
            if not ms.location_match:
                diffs.append(f"lokalizacja: nasza={our_jd.location}, ich={p.location}")
            if not ms.work_mode_match:
                diffs.append(f"tryb pracy: nasz={our_jd.work_mode}, ich={p.work_mode}")
            if ms.salary_delta is not None and abs(ms.salary_delta) > 1000:
                diffs.append(f"delta stawki: {ms.salary_delta:+d} PLN/mies.")
            if diffs:
                _body(doc, "   Roznice vs nasza oferta:")
                for diff in diffs:
                    bp = doc.add_paragraph(style="List Bullet")
                    run2 = bp.add_run(diff)
                    run2.font.name = _FONT_NAME
                    run2.font.size = _FONT_SIZE_BODY


def _add_full_list_section(doc: Any, matched: list[MatchedOffer]) -> None:
    _h2(doc, "3. Pelna lista ofert")

    if not matched:
        _body(doc, "Brak ofert.")
        return

    cols = ["#", "Tytul", "Firma", "Portal", "Score", "Stawka", "Link"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for cell, label in zip(hdr, cols, strict=False):
        cell.text = label
        _set_cell_bg(cell, "1F3A5F")
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = _FONT_NAME
        run.font.size = Pt(9)

    sorted_offers = sorted(
        matched,
        key=lambda m: m.match_score.total if m.match_score else -1,
        reverse=True,
    )

    for i, offer in enumerate(sorted_offers, start=1):
        p = offer.deduped.primary
        ms = offer.match_score
        sal = p.salary
        sal_str = f"{sal.min}-{sal.max} {sal.currency}" if sal else "—"
        score_str = str(ms.total) if ms else "—"
        row = table.add_row().cells
        values = [str(i), p.title, p.company, p.portal, score_str, sal_str, p.url]
        for cell, value in zip(row, values, strict=False):
            cell.text = value
            for run in cell.paragraphs[0].runs:
                run.font.name = _FONT_NAME
                run.font.size = Pt(8)


def _add_footer(doc: Any) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Wygenerowano: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}")
    run.font.name = _FONT_NAME
    run.font.size = Pt(8)
    run.font.color.rgb = _GRAY
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def export_weekly_report(
    matched: list[MatchedOffer],
    our_jd: JDParsed | None,
    week_iso: str,
    output_path: Path,
) -> Path:
    """Generuje tygodniowy raport DOCX.

    Args:
        matched:     Lista MatchedOffer z pipeline.run_matching().
        our_jd:      Sparsowane nasze ogloszenie (None = brak sekcji porownania).
        week_iso:    Numer tygodnia ISO, np. "2026-W20".
        output_path: Sciezka do zapisu pliku .docx.

    Returns:
        output_path po zapisaniu pliku.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = _FONT_NAME
    style.font.size = _FONT_SIZE_BODY

    _h1(doc, f"Recruitment Radar — raport tygodniowy {week_iso}")
    doc.add_paragraph()

    _add_summary_section(doc, matched, our_jd)
    _add_top10_section(doc, matched, our_jd)
    _add_full_list_section(doc, matched)
    _add_footer(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
