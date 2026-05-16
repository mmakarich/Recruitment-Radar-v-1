"""Testy jednostkowe raportu DOCX."""

from __future__ import annotations

import pytest
from docx import Document

from src.export.docx_report import export_weekly_report
from tests.conftest import make_matched_offer


@pytest.fixture
def two_offers(sample_job_offer, sample_salary):
    from datetime import UTC, datetime

    from src.scrapers.base import JobOffer

    offer2 = JobOffer(
        title="Junior React Developer",
        company="Beta Corp",
        portal="nofluffjobs.com",
        url="https://nofluffjobs.com/job/beta-react",
        location="Krakow",
        work_mode="remote",
        seniority="junior",
        tech_stack=("React", "TypeScript"),
        salary=sample_salary,
        published_at=datetime(2026, 5, 12, 10, 0, tzinfo=UTC),
        scraped_at=datetime(2026, 5, 16, 6, 0, tzinfo=UTC),
    )
    return [make_matched_offer(sample_job_offer), make_matched_offer(offer2)]


def test_docx_creates_file(two_offers, tmp_path):
    out = tmp_path / "report.docx"
    result = export_weekly_report(two_offers, None, "2026-W20", out)
    assert result == out
    assert out.exists()
    assert out.stat().st_size > 0


def test_docx_contains_week_in_title(two_offers, tmp_path):
    out = tmp_path / "report.docx"
    export_weekly_report(two_offers, None, "2026-W20", out)
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "2026-W20" in full_text


def test_docx_contains_top10_heading(two_offers, tmp_path):
    out = tmp_path / "report.docx"
    export_weekly_report(two_offers, None, "2026-W20", out)
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "TOP" in full_text


def test_docx_handles_empty_list(tmp_path):
    out = tmp_path / "empty.docx"
    export_weekly_report([], None, "2026-W20", out)
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Brak" in full_text


def test_docx_includes_jd_when_provided(two_offers, tmp_path, sample_jd):
    out = tmp_path / "report.docx"
    export_weekly_report(two_offers, sample_jd, "2026-W20", out)
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    # JDParsed.title powinien pojawic sie w tekscie raportu (sekcja roznic)
    assert sample_jd.title in full_text or "score" in full_text.lower()


def test_docx_footer_has_date(two_offers, tmp_path):
    out = tmp_path / "report.docx"
    export_weekly_report(two_offers, None, "2026-W20", out)
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Wygenerowano" in full_text
